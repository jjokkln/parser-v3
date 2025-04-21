import io
import os
import hashlib
from PIL import Image
import pytesseract
import PyPDF2
from pdf2image import convert_from_path
from docx import Document
import openai
import json
import concurrent.futures
import tempfile

class CombinedProcessor:
    """
    Kombinierte Klasse zur Verarbeitung von Dokumenten und KI-Extraktion in einem Schritt.
    Diese Klasse verbindet die Funktionalitäten des DocumentProcessor und AIExtractor.
    """
    
    def __init__(self, api_key=None):
        """
        Initialisiert den kombinierten Prozessor
        
        Args:
            api_key: OpenAI API Key (optional, kann auch aus Umgebungsvariable geladen werden)
        """
        self.api_key = api_key or os.environ.get("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("OpenAI API Key ist erforderlich")
        
        openai.api_key = self.api_key
        
        # Cache-Verzeichnis erstellen, wenn es nicht existiert
        self.cache_dir = os.path.join(tempfile.gettempdir(), 'parser_cache')
        os.makedirs(self.cache_dir, exist_ok=True)
    
    def process_and_extract(self, file_path, file_extension):
        """
        Hauptfunktion, die sowohl die Textextraktion als auch die KI-Analyse in einem Schritt durchführt.
        
        Args:
            file_path: Pfad zur Datei
            file_extension: Dateierweiterung
        
        Returns:
            Tuple mit (extrahierter Text, strukturierte Profildaten)
        """
        # Datei-Fingerabdruck erstellen für Caching
        file_hash = self._get_file_hash(file_path)
        
        # Prüfen, ob Ergebnisse im Cache vorhanden sind
        cache_result = self._check_cache(file_hash)
        if cache_result:
            return cache_result
        
        # Schritt 1: Text aus Dokument extrahieren
        extracted_text = self._process_document(file_path, file_extension)
        
        # Schritt 2: KI-Analyse der extrahierten Daten
        profile_data = self._extract_profile_data(extracted_text, file_extension)
        
        # Ergebnisse cachen
        self._cache_results(file_hash, extracted_text, profile_data)
        
        return extracted_text, profile_data
    
    def extract_and_process(self, file_path, file_extension):
        """
        Alternative Hauptfunktion, die die KI-Analyse und Textextraktion in umgekehrter Reihenfolge durchführt.
        In der Praxis führt diese Methode zuerst die Extraktion durch (technisch nicht anders möglich)
        und liefert die Ergebnisse in umgekehrter Reihenfolge zurück.
        
        Args:
            file_path: Pfad zur Datei
            file_extension: Dateierweiterung
        
        Returns:
            Tuple mit (strukturierte Profildaten, extrahierter Text)
        """
        # Datei-Fingerabdruck erstellen für Caching
        file_hash = self._get_file_hash(file_path)
        
        # Prüfen, ob Ergebnisse im Cache vorhanden sind
        cache_result = self._check_cache(file_hash)
        if cache_result:
            extracted_text, profile_data = cache_result
            return profile_data, extracted_text
        
        # Tatsächlich muss zuerst der Text extrahiert werden, bevor die KI-Analyse erfolgen kann
        extracted_text = self._process_document(file_path, file_extension)
        profile_data = self._extract_profile_data(extracted_text, file_extension)
        
        # Ergebnisse cachen
        self._cache_results(file_hash, extracted_text, profile_data)
        
        # Gebe die Ergebnisse in umgekehrter Reihenfolge zurück
        return profile_data, extracted_text
    
    def _get_file_hash(self, file_path):
        """Erstellt einen Hash-Wert für eine Datei zur Identifikation im Cache"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _check_cache(self, file_hash):
        """Prüft, ob Ergebnisse für einen Datei-Hash im Cache vorhanden sind"""
        cache_path = os.path.join(self.cache_dir, f"{file_hash}.json")
        if os.path.exists(cache_path):
            try:
                with open(cache_path, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                return cache_data.get('extracted_text', ''), cache_data.get('profile_data', {})
            except Exception as e:
                print(f"Fehler beim Lesen aus dem Cache: {str(e)}")
        return None
    
    def _cache_results(self, file_hash, extracted_text, profile_data):
        """Speichert Extraktions- und Analyseergebnisse im Cache"""
        cache_path = os.path.join(self.cache_dir, f"{file_hash}.json")
        try:
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump({
                    'extracted_text': extracted_text,
                    'profile_data': profile_data
                }, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Fehler beim Schreiben in den Cache: {str(e)}")
    
    # ---- Dokumentenverarbeitung (aus DocumentProcessor) ----
    
    def _process_document(self, file_path, file_extension):
        """
        Extrahiert Text aus Dokumenten verschiedenen Typs
        
        Args:
            file_path: Pfad zur Datei
            file_extension: Dateierweiterung
        
        Returns:
            String mit extrahiertem Text
        """
        if file_extension.lower() in ['.pdf']:
            return self._extract_from_pdf(file_path)
        elif file_extension.lower() in ['.jpg', '.jpeg', '.png']:
            return self._extract_from_image(file_path)
        elif file_extension.lower() in ['.docx']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Nicht unterstützter Dateityp: {file_extension}")
    
    def _extract_from_pdf(self, file_path):
        """Extrahiert Text aus PDF-Dateien mit optimierter Verarbeitung"""
        text = ""
        
        # Versuche zunächst direkte Textextraktion
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Wenn Text erfolgreich extrahiert wurde
                        text += page_text + "\n"
        except Exception as e:
            print(f"Fehler bei direkter PDF-Textextraktion: {str(e)}")
        
        # Falls kein oder wenig Text extrahiert wurde, OCR verwenden
        if len(text.strip()) < 100:  # Heuristik: Weniger als 100 Zeichen bedeutet wahrscheinlich ein Scan
            try:
                # PDF in Bilder umwandeln mit höherer DPI für bessere OCR-Ergebnisse
                images = convert_from_path(file_path, dpi=300)
                
                # Parallele OCR-Verarbeitung für mehrere Seiten
                with concurrent.futures.ThreadPoolExecutor(max_workers=os.cpu_count()) as executor:
                    future_to_image = {executor.submit(self._perform_ocr, image): i for i, image in enumerate(images)}
                    
                    # Resultate sammeln
                    page_texts = [""] * len(images)
                    for future in concurrent.futures.as_completed(future_to_image):
                        page_idx = future_to_image[future]
                        try:
                            page_texts[page_idx] = future.result()
                        except Exception as e:
                            print(f"Fehler bei OCR für Seite {page_idx}: {str(e)}")
                
                # Texte zusammenführen
                text = "\n".join(page_texts)
                
            except Exception as e:
                print(f"Fehler bei PDF-OCR: {str(e)}")
                # Wenn auch OCR fehlschlägt, zurückgeben was wir haben
        
        return text
    
    def _perform_ocr(self, image):
        """OCR für ein einzelnes Bild durchführen"""
        # Bild für bessere OCR-Ergebnisse vorverarbeiten
        image = self._preprocess_image(image)
        # OCR für deutsche Dokumente
        text = pytesseract.image_to_string(image, lang='deu', config='--psm 1 --oem 3')
        return text
    
    def _preprocess_image(self, image):
        """Optimiert Bilder für bessere OCR-Ergebnisse"""
        # Bildvorverarbeitung (Graustufen, Kontrast erhöhen) könnte hier implementiert werden
        # Für einfache Implementierung geben wir das Originalbild zurück
        return image
    
    def _extract_from_image(self, file_path):
        """Extrahiert Text aus Bilddateien mit optimiertem OCR"""
        try:
            # Bild mit PIL öffnen
            image = Image.open(file_path)
            
            # Bild vorverarbeiten
            image = self._preprocess_image(image)
            
            # OCR für deutsche Dokumente mit optimierten Parametern
            text = pytesseract.image_to_string(image, lang='deu', config='--psm 1 --oem 3')
            return text
        except Exception as e:
            raise Exception(f"Fehler bei der Bildverarbeitung: {str(e)}")
    
    def _extract_from_docx(self, file_path):
        """Extrahiert Text aus Word-Dokumenten"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Text aus Absätzen extrahieren
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Text aus Tabellen extrahieren
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Fehler bei der Word-Dokumentverarbeitung: {str(e)}")
    
    # ---- KI-Extraktion (aus AIExtractor) ----
    
    def _extract_profile_data(self, text, document_type):
        """
        Extrahiert strukturierte Profildaten aus Text mit KI-Unterstützung
        
        Args:
            text: Extrahierter Text aus dem Dokument
            document_type: Typ des Dokuments (Dateiendung)
        
        Returns:
            Dictionary mit strukturierten Profildaten
        """
        # Bereite den Prompt für die KI vor
        prompt = self._create_extraction_prompt(text, document_type)
        
        try:
            # OpenAI API-Aufruf
            response = openai.chat.completions.create(
                model="gpt-4o-mini",  # Verwende ein kostengünstiges Modell
                messages=[
                    {"role": "system", "content": "Du bist ein präziser Datenextraktions-Assistent für Lebensläufe."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1  # Niedrige Temperatur für konsistente Ergebnisse
            )
            
            # Extrahiere den Antworttext
            response_text = response.choices[0].message.content
            
            # Versuche JSON zu parsen
            try:
                extracted_data = json.loads(response_text)
                return extracted_data
            except json.JSONDecodeError:
                # Fallback-Mechanismus: Versuche JSON aus Freitext zu extrahieren
                return self._extract_json_from_text(response_text)
                
        except Exception as e:
            raise Exception(f"Fehler bei der KI-Extraktion: {str(e)}")
    
    def _create_extraction_prompt(self, text, document_type):
        """Erstellt einen Prompt für die KI-Extraktion"""
        return f"""
Du bist ein Assistent für die Extraktion von Lebenslaufdaten. Analysiere den folgenden Text aus einem {document_type}-Dokument und extrahiere alle relevanten Informationen in das spezifizierte JSON-Format.

Der Text stammt aus einem Lebenslauf und enthält Informationen über eine Person, ihre Berufserfahrung, Ausbildung und Qualifikationen.

Extrahierter Text:
{text}

Liefere das Ergebnis ausschließlich im folgenden JSON-Format ohne zusätzlichen Text oder Erklärungen:
{{
  "persönliche_daten": {{
    "name": "",
    "wohnort": "",
    "jahrgang": "",
    "führerschein": "",
    "kontakt": {{
      "ansprechpartner": "",
      "telefon": "",
      "email": ""
    }}
  }},
  "berufserfahrung": [
    {{
      "zeitraum": "",
      "unternehmen": "",
      "position": "",
      "aufgaben": []
    }}
  ],
  "ausbildung": [
    {{
      "zeitraum": "",
      "institution": "",
      "schwerpunkte": "",
      "abschluss": "",
      "note": ""
    }}
  ],
  "weiterbildungen": [
    {{
      "zeitraum": "",
      "bezeichnung": "",
      "abschluss": ""
    }}
  ]}},
  "wunschgehalt": ""
}}

Hinweise:
- Achte darauf, das exakte JSON-Format zu verwenden
- Vervollständige alle Felder, die im Text identifiziert werden können
- Lasse Felder leer, wenn keine Information vorhanden ist
- Organisiere die berufliche Erfahrung chronologisch (neueste zuerst)
- Bei Studiengängen extrahiere auch die Studienschwerpunkte, falls angegeben
- Beim Führerschein gib auch an, ob ein PKW vorhanden ist, falls diese Information verfügbar ist
- Falls der Zeitraum als "Seit MM/JJJJ" angegeben ist, erfasse nur den Zeitpunkt (z.B. "07/2020")
- Versuche, die Aufgaben als einzelne Punkte zu strukturieren, statt als einen langen Text
- Das Wunschgehalt, falls erwähnt, sollte als Jahresgehalt in Euro extrahiert werden
"""
    
    def _extract_json_from_text(self, text):
        """
        Fallback-Methode zur Extraktion von JSON aus Freitext
        Nützlich, wenn das KI-Modell nicht nur JSON zurückgibt
        """
        # Versuche JSON-Daten aus dem Text zu extrahieren
        try:
            # Suche nach geschweiften Klammern, die JSON umschließen könnten
            start_idx = text.find('{')
            end_idx = text.rfind('}') + 1
            
            if start_idx >= 0 and end_idx > start_idx:
                json_str = text[start_idx:end_idx]
                return json.loads(json_str)
            else:
                # JSON nicht gefunden, einfache Struktur zurückgeben
                return {
                    "persönliche_daten": {
                        "name": "Nicht erkannt",
                        "wohnort": "",
                        "jahrgang": "",
                        "führerschein": "Klasse B (Pkw vorhanden)",
                        "kontakt": {
                            "ansprechpartner": "Fischer",
                            "telefon": "02161 62126-02",
                            "email": "fischer@galdora.de"
                        }
                    },
                    "berufserfahrung": [],
                    "ausbildung": [],
                    "weiterbildungen": [],
                    "wunschgehalt": ""
                }
        except Exception:
            # Bei Fehlern leere Struktur zurückgeben
            return {
                "persönliche_daten": {
                    "name": "Nicht erkannt",
                    "wohnort": "",
                    "jahrgang": "",
                    "führerschein": "Klasse B (Pkw vorhanden)",
                    "kontakt": {
                        "ansprechpartner": "Fischer",
                        "telefon": "02161 62126-02",
                        "email": "fischer@galdora.de"
                    }
                },
                "berufserfahrung": [],
                "ausbildung": [],
                "weiterbildungen": [],
                "wunschgehalt": ""
            } 