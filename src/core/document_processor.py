import io
import os
from PIL import Image
import pytesseract
import PyPDF2
from pdf2image import convert_from_path
from docx import Document

class DocumentProcessor:
    """Klasse zur Verarbeitung verschiedener Dokumenttypen"""
    
    def process_document(self, file_path, file_extension):
        """
        Extrahiert Text aus Dokumenten verschiedenen Typs
        
        Args:
            file_path: Pfad zur Datei
            file_extension: Dateierweiterung
        
        Returns:
            String mit extrahiertem Text
        """
        if file_extension.lower() in ['.pdf']:
            return self._extract_from_pdf(file_path)
        elif file_extension.lower() in ['.jpg', '.jpeg', '.png']:
            return self._extract_from_image(file_path)
        elif file_extension.lower() in ['.docx']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Nicht unterstützter Dateityp: {file_extension}")
    
    def _extract_from_pdf(self, file_path):
        """Extrahiert Text aus PDF-Dateien"""
        text = ""
        
        # Versuche zunächst direkte Textextraktion
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Wenn Text erfolgreich extrahiert wurde
                        text += page_text + "\n"
        except Exception as e:
            print(f"Fehler bei direkter PDF-Textextraktion: {str(e)}")
        
        # Falls kein oder wenig Text extrahiert wurde, OCR verwenden
        if len(text.strip()) < 100:  # Heuristik: Weniger als 100 Zeichen bedeutet wahrscheinlich ein Scan
            try:
                # PDF in Bilder umwandeln und OCR durchführen
                images = convert_from_path(file_path)
                for i, image in enumerate(images):
                    page_text = pytesseract.image_to_string(image, lang='deu')  # Deutsch für deutsche Dokumente
                    text += page_text + "\n"
            except Exception as e:
                print(f"Fehler bei PDF-OCR: {str(e)}")
                # Wenn auch OCR fehlschlägt, zurückgeben was wir haben
        
        return text
    
    def _extract_from_image(self, file_path):
        """Extrahiert Text aus Bilddateien mit OCR"""
        try:
            image = Image.open(file_path)
            # OCR für deutsche Dokumente (für andere Sprachen anpassen)
            text = pytesseract.image_to_string(image, lang='deu')
            return text
        except Exception as e:
            raise Exception(f"Fehler bei der Bildverarbeitung: {str(e)}")
    
    def _extract_from_docx(self, file_path):
        """Extrahiert Text aus Word-Dokumenten"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Text aus Absätzen extrahieren
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Text aus Tabellen extrahieren
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Fehler bei der Word-Dokumentverarbeitung: {str(e)}")
