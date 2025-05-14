from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, HRFlowable, PageBreak, KeepTogether, FrameBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import PageTemplate
import io
import os
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.utils.image_utils import get_image_path, ensure_images_in_static

class ProfileGenerator:
    """Klasse zur Erstellung von PDF-Profilen aus strukturierten Daten"""
    
    def __init__(self):
        """Initialisiert den Profil-Generator mit Stilen"""
        self.styles = getSampleStyleSheet()
        self.custom_styles = self._create_custom_styles()
    
    def generate_profile(self, profile_data, output_path, template="professional", format="pdf"):
        """
        Generiert ein Profil aus den extrahierten Daten im gewünschten Format (PDF oder DOCX)
        
        Args:
            profile_data: Dictionary mit Profildaten
            output_path: Pfad für die Ausgabedatei
            template: Art der Vorlage (professional, classic, modern, minimalist)
            format: Format der Ausgabedatei ("pdf" oder "docx")
        
        Returns:
            Pfad zur generierten Datei
        """
        # Prüfe, ob profile_data ein gültiges Dictionary ist
        if not isinstance(profile_data, dict):
            print(f"Warnung: profile_data ist kein Dictionary. Typ: {type(profile_data)}")
            profile_data = {}
        
        # Überprüfe, ob der Ausgabepfad gültig ist
        if not output_path:
            raise ValueError("Der Ausgabepfad darf nicht leer sein.")
            
        # Überprüfe, ob das Verzeichnis existiert, falls nicht, erstelle es
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Je nach Format unterschiedliche Generierungsmethode aufrufen
        if format.lower() == "pdf":
            return self._generate_pdf_profile(profile_data, output_path, template)
        elif format.lower() == "docx":
            # Sicherstellen, dass der Ausgabepfad auf .docx endet
            if not output_path.lower().endswith('.docx'):
                base_path = os.path.splitext(output_path)[0]
                output_path = base_path + '.docx'
            return self._generate_docx_profile(profile_data, output_path, template)
        else:
            raise ValueError(f"Ungültiges Format: {format}. Unterstützte Formate: 'pdf', 'docx'")
    
    def _generate_pdf_profile(self, profile_data, output_path, template="professional"):
        """
        Generiert ein PDF-Profil aus den extrahierten Daten
        
        Args:
            profile_data: Dictionary mit Profildaten
            output_path: Pfad für die Ausgabedatei
            template: Art der Vorlage (professional, classic, modern, minimalist)
        
        Returns:
            Pfad zur generierten PDF-Datei
        """
        try:
            # Footer text
            footer_text = "GALDORA Personalmanagement GmbH Co.KG\nVolksgartenstr. 85-89, 41065 Mönchengladbach\nE-Mail: info@galdora.de / Web: www.galdora.de"
            
            # Erstelle ein Template mit Hauptframe und Footerframe
            def add_page_number(canvas, doc):
                # Footer
                canvas.saveState()
                footer_style = ParagraphStyle(
                    'Footer',
                    parent=self.styles['Normal'],
                    fontSize=7,
                    fontName='Helvetica',
                    alignment=1,  # Zentriert
                    textColor=colors.black
                )
                p = Paragraph(footer_text, footer_style)
                w, h = p.wrap(doc.width, doc.bottomMargin)
                p.drawOn(canvas, doc.leftMargin, 15*mm)
                canvas.restoreState()
            
            # Erstelle das PDF-Dokument mit dem Footer-Callback
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=20*mm, 
                leftMargin=20*mm,
                topMargin=20*mm, 
                bottomMargin=25*mm  # Etwas mehr Platz für Footer
            )
            
            # Erstelle die Dokumentelemente basierend auf der gewählten Vorlage
            elements = self._create_document_elements(profile_data, template)
            
            # Erstelle das PDF mit dem Footer auf jeder Seite
            doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
            
            return output_path
        except Exception as e:
            print(f"Fehler bei der PDF-Generierung: {str(e)}")
            raise
    
    def _generate_docx_profile(self, profile_data, output_path, template="professional"):
        """
        Generiert ein Word-Dokument (DOCX) aus den extrahierten Daten
        
        Args:
            profile_data: Dictionary mit Profildaten
            output_path: Pfad für die Ausgabedatei
            template: Art der Vorlage (professional, classic, modern, minimalist)
        
        Returns:
            Pfad zur generierten DOCX-Datei
        """
        try:
            # Erstelle ein neues Word-Dokument
            doc = docx.Document()
            
            # Seiteneinstellungen (A4)
            section = doc.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            
            # Definiere Stile für das Word-Dokument
            # Header-Style für GALDORA Logo
            logo_style = doc.styles.add_style('GaldoraLogo', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            logo_style.font.size = Pt(36)
            logo_style.font.bold = True
            logo_style.font.color.rgb = RGBColor(0, 0, 0) # Schwarz
            
            # Erstelle Italic Style
            italic_style = doc.styles.add_style('ItalicStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            italic_style.font.size = Pt(10)
            italic_style.font.italic = True
            
            # Korrekte Pfade zum Verzeichnis der Quellendateien
            current_dir = os.path.dirname(os.path.abspath(__file__))
            sources_dir = os.path.join(os.path.dirname(os.path.dirname(current_dir)), 'sources')
            
            # Wenn sources nicht existiert, versuche relative Pfade vom Arbeitsverzeichnis
            if not os.path.exists(sources_dir):
                sources_dir = 'sources'
                # Erstelle das Verzeichnis, falls es nicht existiert
                if not os.path.exists(sources_dir):
                    os.makedirs(sources_dir, exist_ok=True)
                    print(f"Verzeichnis '{sources_dir}' wurde erstellt")
            
            # Logo einfügen, wenn verfügbar
            logo_path = os.path.join(sources_dir, 'Galdoralogo.png')
            
            if os.path.exists(logo_path) and os.path.isfile(logo_path):
                # Paragraph für das Logo erstellen
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Cm(5))  # Logo mit 5cm Breite einfügen
            else:
                # Fallback, wenn kein Logo gefunden wurde
                header = doc.add_paragraph("GALDORA", style='GaldoraLogo')
                header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Tagline
            tagline = doc.add_paragraph("Wir verbinden Menschen und Technologie", style='ItalicStyle')
            tagline.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Füge Leerzeile hinzu
            doc.add_paragraph()
            
            # Überschrift "Profil"
            title_style = doc.styles.add_style('ProfilTitle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            # GALDORA Blau (ca. #1973B8)
            title_style.font.color.rgb = RGBColor(25, 115, 184)
            
            doc.add_paragraph("Profil", style='ProfilTitle')
            
            # Persönliche Daten
            personal_data = profile_data.get("persönliche_daten", {})
            name = personal_data.get("name", "")
            doc.add_paragraph(name).bold = True
            
            # Erstelle eine 2-spaltige Tabelle für persönliche Daten
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            
            # Wohnort
            if personal_data.get("wohnort"):
                row_cells = table.add_row().cells
                row_cells[0].text = "Wohnort:"
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].text = personal_data.get("wohnort", "")
            
            # Jahrgang
            if personal_data.get("jahrgang"):
                row_cells = table.add_row().cells
                row_cells[0].text = "Jahrgang:"
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].text = personal_data.get("jahrgang", "")
            
            # Führerschein
            if personal_data.get("führerschein"):
                row_cells = table.add_row().cells
                row_cells[0].text = "Führerschein:"
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].text = personal_data.get("führerschein", "")
            
            # Berufserfahrung
            doc.add_paragraph().add_run().add_break()  # Leerzeile
            heading_style = doc.styles.add_style('Heading2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True
            
            doc.add_paragraph("Beruflicher Werdegang", style='Heading2')
            
            # Füge Berufserfahrung hinzu
            for experience in profile_data.get("berufserfahrung", []):
                # Tabelle für jede Berufserfahrung
                exp_table = doc.add_table(rows=0, cols=2)
                exp_table.style = 'Table Grid'
                exp_table.autofit = True
                
                # Zeitraum
                row_cells = exp_table.add_row().cells
                row_cells[0].text = experience.get("zeitraum", "")
                row_cells[0].paragraphs[0].runs[0].bold = True
                
                # Firma und Position
                row_cells[1].text = experience.get("firma", "")
                row_cells[1].paragraphs[0].runs[0].bold = True
                row_cells[1].add_paragraph(experience.get("position", "")).italic = True
                
                # Beschreibung (wenn vorhanden)
                if experience.get("beschreibung"):
                    row_cells = exp_table.add_row().cells
                    row_cells[0].merge(row_cells[1])
                    row_cells[0].text = experience.get("beschreibung", "")
                
                doc.add_paragraph()  # Leerzeile
            
            # Ausbildung
            doc.add_paragraph("Ausbildung", style='Heading2')
            
            # Füge Ausbildung hinzu
            for education in profile_data.get("ausbildung", []):
                # Tabelle für jede Ausbildung
                edu_table = doc.add_table(rows=0, cols=2)
                edu_table.style = 'Table Grid'
                edu_table.autofit = True
                
                # Zeitraum
                row_cells = edu_table.add_row().cells
                row_cells[0].text = education.get("zeitraum", "")
                row_cells[0].paragraphs[0].runs[0].bold = True
                
                # Institution und Abschluss
                row_cells[1].text = education.get("institution", "")
                row_cells[1].paragraphs[0].runs[0].bold = True
                row_cells[1].add_paragraph(education.get("abschluss", "")).italic = True
                
                doc.add_paragraph()  # Leerzeile
            
            # Trennlinie zwischen Ausbildung und Weiterbildungen
            doc.add_paragraph()
            doc.add_paragraph("Weiterbildungen", style='Heading2')
            
            # Weiterbildungen
            if profile_data.get("weiterbildungen"):
                for training in profile_data.get("weiterbildungen", []):
                    # Tabelle für jede Weiterbildung
                    train_table = doc.add_table(rows=0, cols=2)
                    train_table.style = 'Table Grid'
                    train_table.autofit = True
                    
                    # Zeitraum
                    row_cells = train_table.add_row().cells
                    row_cells[0].text = training.get("zeitraum", "")
                    row_cells[0].paragraphs[0].runs[0].bold = True
                    
                    # Bezeichnung und Abschluss
                    row_cells[1].text = training.get("bezeichnung", "")
                    row_cells[1].paragraphs[0].runs[0].bold = True
                    
                    if training.get("abschluss"):
                        row_cells[1].add_paragraph(training.get("abschluss", "")).italic = True
                    
                    doc.add_paragraph()  # Leerzeile
            else:
                doc.add_paragraph("Keine Weiterbildungen angegeben", style='Normal')
            
            # Erstelle Seitenumbruch vor Ansprechpartner-Seite
            doc.add_page_break()
            
            # Ansprechpartner-Seite
            # Logo einfügen, wenn verfügbar
            if os.path.exists(logo_path) and os.path.isfile(logo_path):
                # Paragraph für das Logo erstellen
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Cm(5))  # Logo mit 5cm Breite einfügen
            else:
                # Fallback, wenn kein Logo gefunden wurde
                header = doc.add_paragraph("GALDORA", style='GaldoraLogo')
                header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Tagline
            tagline = doc.add_paragraph("Wir verbinden Menschen und Technologie", style='ItalicStyle')
            tagline.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Leerzeile
            doc.add_paragraph()
            
            # Ansprechpartner
            doc.add_paragraph("Kontakt", style='ProfilTitle')
            
            # Kontaktdaten
            contact_data = personal_data.get("kontakt", {})
            ansprechpartner = contact_data.get("ansprechpartner", "")
            
            # Bestimme die korrekte Anrede (Ausnahmebehandlung für Melike Demirkol)
            if ansprechpartner == "Melike Demirkol":
                anrede = f"Frau Demirkol"
            else:
                nachname = ansprechpartner.split()[-1] if ansprechpartner else 'Fischer'
                anrede = f"Herr {nachname}"
            
            telefon = contact_data.get("telefon", "")
            email = contact_data.get("email", "")
            
            if ansprechpartner:
                doc.add_paragraph("IHR ANSPRECHPARTNER").bold = True
                contact_para = doc.add_paragraph(ansprechpartner)
                contact_para.paragraph_format.space_after = Pt(5)
                
                # Tel und E-Mail
                if telefon:
                    tel_para = doc.add_paragraph(f"Tel.: {telefon}")
                    tel_para.paragraph_format.space_after = Pt(5)
                
                if email:
                    email_para = doc.add_paragraph(f"E-Mail: {email}")
                    email_para.paragraph_format.space_after = Pt(5)
            
            # Wunschgehalt
            wunschgehalt = profile_data.get("wunschgehalt", "")
            if wunschgehalt:
                doc.add_paragraph()  # Leerzeile
                doc.add_paragraph("INFORMATIONEN ZUR BEWERBUNG").bold = True
                gehalt_para = doc.add_paragraph(f"Gehaltsvorstellung: {wunschgehalt}")
                gehalt_para.paragraph_format.space_after = Pt(5)
                
                # Verfügbarkeit anzeigen
                verfuegbarkeit_status = profile_data.get("verfuegbarkeit_status", "")
                if verfuegbarkeit_status:
                    verf_para = doc.add_paragraph(f"Verfügbarkeit: {verfuegbarkeit_status}")
                    verf_para.paragraph_format.space_after = Pt(5)
                    
                    # Details zur Verfügbarkeit anzeigen, wenn vorhanden
                    verfuegbarkeit_details = profile_data.get("verfuegbarkeit_details", "")
                    if verfuegbarkeit_details:
                        details_para = doc.add_paragraph(f"Details zur Verfügbarkeit: {verfuegbarkeit_details}")
                        details_para.paragraph_format.space_after = Pt(5)
            # Nur Verfügbarkeit anzeigen, wenn kein Wunschgehalt vorhanden ist
            elif profile_data.get("verfuegbarkeit_status", ""):
                doc.add_paragraph()  # Leerzeile
                doc.add_paragraph("INFORMATIONEN ZUR BEWERBUNG").bold = True
                
                verfuegbarkeit_status = profile_data.get("verfuegbarkeit_status", "")
                verf_para = doc.add_paragraph(f"Verfügbarkeit: {verfuegbarkeit_status}")
                verf_para.paragraph_format.space_after = Pt(5)
                
                verfuegbarkeit_details = profile_data.get("verfuegbarkeit_details", "")
                if verfuegbarkeit_details:
                    details_para = doc.add_paragraph(f"Details zur Verfügbarkeit: {verfuegbarkeit_details}")
                    details_para.paragraph_format.space_after = Pt(5)
            
            # Speichere das Dokument
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"Fehler bei der DOCX-Generierung: {str(e)}")
            raise
    
    def _create_custom_styles(self):
        """Erstellt benutzerdefinierte Stile für das Dokument"""
        custom_styles = {}
        
        # GALDORA Überschrift-Stil (große schwarze Schrift)
        custom_styles['GaldoraLogo'] = ParagraphStyle(
            'GaldoraLogo',
            parent=self.styles['Normal'],
            fontSize=36,
            fontName='Helvetica-Bold',
            textColor=colors.black,
            spaceAfter=0.1*cm
        )

        # Tagline unter dem Logo
        custom_styles['Tagline'] = ParagraphStyle(
            'Tagline',
            parent=self.styles['Italic'],
            fontSize=10,
            fontName='Helvetica-Oblique',
            textColor=colors.black,
            spaceAfter=1.5*cm
        )
        
        # Profil Überschrift - nach links ausgerichtet statt zentriert
        custom_styles['ProfilTitle'] = ParagraphStyle(
            'ProfilTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#1973B8'),  # GALDORA Blau
            spaceBefore=0.3*cm,
            spaceAfter=0.5*cm,
            alignment=0,  # Links ausgerichtet (0) statt zentriert (1)
            underline=0  # Keine Unterstreichung
        )
        
        # Name - nach links ausgerichtet statt zentriert
        custom_styles['Name'] = ParagraphStyle(
            'Name',
            parent=self.styles['Normal'],
            fontSize=14,
            fontName='Helvetica',
            spaceBefore=0.2*cm,
            spaceAfter=0.8*cm,
            alignment=0  # Links ausgerichtet (0) statt zentriert (1)
        )
        
        # Überschrift für Abschnitte (Beruflicher Werdegang, Ausbildung, etc.)
        custom_styles['Heading2'] = ParagraphStyle(
            'Heading2',
            parent=self.styles['Heading2'],
            fontSize=12,
            fontName='Helvetica-Bold',
            textColor=colors.black,
            spaceBefore=0.7*cm,
            spaceAfter=0.3*cm,
            underline=0  # Keine Unterstreichung
        )
        
        # Ansprechpartner Überschrift - ALLES GROß
        custom_styles['ContactHeader'] = ParagraphStyle(
            'ContactHeader',
            parent=self.styles['Normal'],
            fontSize=9,
            fontName='Helvetica-Bold',
            spaceBefore=0.2*cm,
            spaceAfter=0.1*cm,
            underline=0
        )
        
        # Ansprechpartner Daten (grau)
        custom_styles['ContactData'] = ParagraphStyle(
            'ContactData',
            parent=self.styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            spaceAfter=0.1*cm,
            leftIndent=0*cm
        )
        
        # Normaler Text
        custom_styles['Normal'] = ParagraphStyle(
            'Normal',
            parent=self.styles['Normal'],
            fontSize=9,
            spaceAfter=0.1*cm
        )
        
        # Label-Text (fett) mit Text in gleicher Zeile
        custom_styles['LabelInline'] = ParagraphStyle(
            'LabelInline',
            parent=self.styles['Normal'],
            fontSize=9,
            fontName='Helvetica-Bold',
            spaceAfter=0.1*cm
        )
        
        # Label-Text (fett)
        custom_styles['Label'] = ParagraphStyle(
            'Label',
            parent=self.styles['Normal'],
            fontSize=9,
            fontName='Helvetica-Bold',
            spaceAfter=0.1*cm
        )
        
        # Firmenname - Hervorgehoben
        custom_styles['Company'] = ParagraphStyle(
            'Company',
            parent=self.styles['Normal'],
            fontSize=9,
            fontName='Helvetica-Bold',
            spaceAfter=0.1*cm
        )
        
        # Position/Rolle - Kursiv
        custom_styles['Position'] = ParagraphStyle(
            'Position',
            parent=self.styles['Italic'],
            fontSize=9,
            fontName='Helvetica-Oblique',
            spaceAfter=0.2*cm
        )
        
        # Zeitraum - Linksbündig für erste Spalte 
        custom_styles['Period'] = ParagraphStyle(
            'Period',
            parent=self.styles['Normal'],
            fontSize=9,
            fontName='Helvetica-Bold',
            alignment=0,  # Linksbündig
            spaceAfter=0.1*cm
        )
        
        # Aufgabenpunkt - für rechte Spalte
        custom_styles['TaskPoint'] = ParagraphStyle(
            'TaskPoint',
            parent=self.styles['Normal'],
            fontSize=9,
            alignment=0,  # Linksbündig
            leftIndent=0.5*cm,
            bulletIndent=0.3*cm,
            spaceAfter=0.1*cm
        )
        
        # GALDORA Fußzeile
        custom_styles['Footer'] = ParagraphStyle(
            'Footer',
            parent=self.styles['Normal'],
            fontSize=7,
            fontName='Helvetica',
            alignment=1,  # Zentriert
            spaceBefore=0.2*cm,
            textColor=colors.black
        )
        
        return custom_styles
    
    def _create_document_elements(self, profile_data, template="professional"):
        
        """
        Erstellt die Elemente für das PDF-Dokument basierend auf dem Design der Profilvorlage
        
        Args:
            profile_data: Dictionary mit Profildaten
            template: Art der Vorlage (professional, classic, modern, minimalist)
        
        Returns:
            Liste mit PDF-Elementen
        """
        elements = []
        
        try:
            # Stelle sicher, dass alle Bilder im static-Verzeichnis verfügbar sind (für HTTPS-Kompatibilität)
            ensure_images_in_static()
            
            # Verwende die Image-Utility für den Pfad
            # Bei HTTPS-Server nutze use_static=True, für Localhost use_static=False
            # Standardmäßig verwenden wir das static-Verzeichnis für HTTPS-Kompatibilität
            use_https_compatible = True  # Explizit auf True gesetzt für HTTPS-Kompatibilität
            
            # Je nach gewähltem Template unterschiedliche Layouts erstellen
            if template == "modern":
                # MODERN TEMPLATE (basierend auf dem Bild)
                # Ensure we have personal_data initialized at the start
                personal_data = profile_data.get('persönliche_daten', {})
                
                # Verwende das GALDORA Logo aus dem richtigen Ordner
                logo_path = get_image_path('Galdoralogo.png', use_static=use_https_compatible)
                
                # Erstelle eine Tabelle für das Logo oben
                if os.path.exists(logo_path) and os.path.isfile(logo_path):
                    try:
                        # Logo-Größe korrigieren (Original-Proportionen beibehalten)
                        from PIL import Image as PILImage
                        img_pil = PILImage.open(logo_path)
                        img_width, img_height = img_pil.size
                        aspect_ratio = img_width / img_height
                        
                        target_width = 150
                        target_height = target_width / aspect_ratio
                        
                        img = Image(logo_path, width=target_width, height=target_height)
                        elements.append(Table([[img]], colWidths=[A4[0] - 40*mm]))
                        elements.append(Spacer(1, 0.5*cm))
                    except Exception as e:
                        print(f"Fehler beim Laden des Logos: {str(e)}")
                        elements.append(Paragraph("GALDORA", self.custom_styles['GaldoraLogo']))
                else:
                    print(f"Logo-Datei nicht gefunden: {logo_path}")
                    elements.append(Paragraph("GALDORA", self.custom_styles['GaldoraLogo']))
                
                # Persönliche Daten bereits oben initialisiert
                name = personal_data.get('name', 'Profil')
                
                # Moderne Layout mit zwei Spalten erstellen
                # Linke Spalte: Foto und persönliche Daten
                # Rechte Spalte: Berufserfahrung, Ausbildung, Skills, etc.
                
                # Erstelle Stil für den Namen im modernen Template (größer und fett)
                modern_name_style = ParagraphStyle(
                    'ModernName',
                    parent=self.custom_styles['Normal'],
                    fontSize=16,
                    fontName='Helvetica-Bold',
                    spaceAfter=0.5*cm,
                    alignment=0
                )
                
                # Erstelle Stil für die Kategorie-Überschriften (links)
                modern_category_style = ParagraphStyle(
                    'ModernCategory',
                    parent=self.custom_styles['Normal'],
                    fontSize=14,
                    fontName='Helvetica-Bold',
                    textColor=colors.HexColor('#1973B8'),  # GALDORA Blau
                    spaceAfter=0.3*cm,
                    alignment=0
                )
                
                # Erstelle Stil für die linke Spalte (Kategorien)
                modern_left_style = ParagraphStyle(
                    'ModernLeft',
                    parent=self.custom_styles['Normal'],
                    fontSize=10,
                    fontName='Helvetica-Bold',
                    spaceAfter=0.2*cm,
                    alignment=0
                )
                
                # Erstelle Stil für die rechte Spalte (Content)
                modern_right_style = ParagraphStyle(
                    'ModernRight', 
                    parent=self.custom_styles['Normal'],
                    fontSize=10,
                    spaceAfter=0.2*cm,
                    alignment=0
                )
                
                # Inhalte für beide Spalten vorbereiten
                left_content = []
                right_content = []
                
                # Name oben rechts platzieren
                right_content.append(Paragraph(name, modern_name_style))
                
                # Linke Spalte: Profilbild
                # Überprüfen, ob ein Profilbild vorhanden ist
                profile_image_path = personal_data.get("profile_image", None)
                
                if profile_image_path and os.path.exists(profile_image_path):
                    try:
                        # Profilbild laden und anzeigen
                        from PIL import Image as PILImage
                        img_pil = PILImage.open(profile_image_path)
                        img_width, img_height = img_pil.size
                        aspect_ratio = img_width / img_height
                        
                        # Standardgröße für das Profilbild
                        photo_width = 120
                        photo_height = photo_width / aspect_ratio
                        
                        # Profilbild im PDF einfügen
                        profile_img = Image(profile_image_path, width=photo_width, height=photo_height)
                        left_content.append(profile_img)
                    except Exception as e:
                        print(f"Fehler beim Laden des Profilbilds: {str(e)}")
                        # Fallback: "FOTO" Platzhalter anzeigen
                        photo_placeholder = Paragraph("FOTO", self.custom_styles['Normal'])
                        left_content.append(photo_placeholder)
                else:
                    # Kein Bild vorhanden, Platzhalter anzeigen
                    photo_placeholder = Paragraph("FOTO", self.custom_styles['Normal'])
                    left_content.append(photo_placeholder)
                left_content.append(Spacer(1, 1*cm))
            
                # Linke Spalte: SOFTWARE
                left_content.append(Paragraph("SOFTWARE", modern_category_style))
                skills = profile_data.get('skills', {})
                software_skills = skills.get('software', [])
                for skill in software_skills:
                    try:
                        name = skill.get('name', '')
                        level = skill.get('level', '')
                        if name and level:
                            left_content.append(Paragraph(f"• {name} - {level}", self.custom_styles['Normal']))
                        elif name:
                            left_content.append(Paragraph(f"• {name}", self.custom_styles['Normal']))
                    except Exception as e:
                        print(f"Fehler bei Software-Skill: {str(e)}")
                left_content.append(Spacer(1, 1*cm))
                
                # Linke Spalte: SPRACHEN
                left_content.append(Paragraph("SPRACHEN", modern_category_style))
                language_skills = skills.get('sprachen', [])
                for skill in language_skills:
                    try:
                        name = skill.get('name', '')
                        level = skill.get('level', '')
                        if name and level:
                            if name == "Deutsch":
                                left_content.append(Paragraph(f"{name}: Muttersprache", self.custom_styles['Normal']))
                            else:
                                # Sprachlevel als Punkte darstellen
                                level_dots = ""
                                if level.lower() in ["gut", "sehr gut", "fließend", "c1", "c2"]:
                                    level_dots = "●●●●○"
                                elif level.lower() in ["mittel", "befriedigend", "b1", "b2"]:
                                    level_dots = "●●●○○"
                                else:
                                    level_dots = "●●○○○"
                                    
                                left_content.append(Paragraph(f"{name}", self.custom_styles['Normal']))
                                left_content.append(Paragraph(f"{level_dots}", self.custom_styles['Normal']))
                        elif name:
                            left_content.append(Paragraph(f"{name}", self.custom_styles['Normal']))
                    except Exception as e:
                        print(f"Fehler bei Sprach-Skill: {str(e)}")
                left_content.append(Spacer(1, 1*cm))
                
                # Linke Spalte: HOBBYS
                left_content.append(Paragraph("HOBBYS", modern_category_style))
                hobbys = profile_data.get('hobbys', [])
                if hobbys:
                    for hobby in hobbys:
                        left_content.append(Paragraph(f"• {hobby}", self.custom_styles['Normal']))
                else:
                    left_content.append(Paragraph("• Keine Angaben", self.custom_styles['Normal']))
                
                # Rechte Spalte: KURZPROFIL
                right_content.append(Paragraph("KURZPROFIL", modern_category_style))
                kurzprofil = profile_data.get('kurzprofil', '')
                if kurzprofil:
                    right_content.append(Paragraph(kurzprofil, self.custom_styles['Normal']))
                else:
                    right_content.append(Paragraph("Kein Kurzprofil angegeben", self.custom_styles['Normal']))
                right_content.append(Spacer(1, 0.5*cm))
                
                # Rechte Spalte: BERUFSERFAHRUNG
                right_content.append(Paragraph("BERUFSERFAHRUNG", modern_category_style))
                berufserfahrung = profile_data.get('berufserfahrung', [])
                if berufserfahrung:
                    for erfahrung in berufserfahrung:
                        try:
                            zeitraum = erfahrung.get('zeitraum', '')
                            unternehmen = erfahrung.get('unternehmen', '')
                            position = erfahrung.get('position', '')
                            
                            # Zeitraum und Unternehmen
                            right_column_content = [
                                Paragraph(unternehmen, self.custom_styles['Company']),
                                Paragraph(position, self.custom_styles['Position'])
                            ]
                            
                            # Aufgaben auf maximal 4 begrenzen
                            aufgaben = erfahrung.get('aufgaben', [])
                            aufgaben_formatted = []
                            for i, aufgabe in enumerate(aufgaben[:4]):  # Maximal 4 Aufgaben
                                aufgaben_formatted.append(Paragraph(f"• {aufgabe}", self.custom_styles['Normal']))
                            
                            # Aufgaben zur rechten Spalte hinzufügen
                            right_column_content.extend(aufgaben_formatted)
                            
                            # Bereite die Tabellendaten vor
                            data = [[Paragraph(zeitraum, self.custom_styles['Period']), right_column_content[0]]]
                            
                            # Füge weitere Zeilen hinzu
                            for i in range(1, len(right_column_content)):
                                data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[i]])
                            
                            # Tabelle mit definierter Breite (10% links, 75% rechts)
                            col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                            
                            table = Table(data, colWidths=col_widths)
                            table.setStyle(TableStyle([
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                                ('LEFTPADDING', (1, 0), (1, -1), 2*cm),  # Ca. 2cm Einrückung für bessere Lesbarkeit
                            ]))
                            
                            # Wir verpacken die Tabelle und den Spacer in KeepTogether, damit sie nicht über eine Seite verteilt werden
                            entry_elements = [table, Spacer(1, 0.3*cm)]
                            elements.append(KeepTogether(entry_elements))
                        except Exception as e:
                            print(f"Fehler bei der Verarbeitung einer Berufserfahrung: {str(e)}")
                            # Einfache Darstellung als Fallback
                            elements.append(Paragraph(f"{zeitraum} - {unternehmen} - {position}", self.custom_styles['Normal']))
                            elements.append(Spacer(1, 0.3*cm))
                else:
                    elements.append(Paragraph("Keine Berufserfahrung angegeben", self.custom_styles['Normal']))
                
                # Ausbildung
                elements.append(Spacer(1, 0.5*cm))
                # Trennlinie zwischen Berufserfahrung und Ausbildung
                elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.lightgrey))
                elements.append(Spacer(1, 0.5*cm))
                elements.append(Paragraph("Ausbildung/ Weiterbildung", self.custom_styles['Heading2']))
                
                # Ausbildungen im gleichen Format wie Berufserfahrung darstellen
                ausbildungen = profile_data.get('ausbildung', [])
                if ausbildungen:
                    for ausbildung in ausbildungen:
                        try:
                            # Zeitraum
                            zeitraum = ausbildung.get('zeitraum', '')
                            
                            # Institution und Abschluss
                            institution = ausbildung.get('institution', '')
                            abschluss = ausbildung.get('abschluss', '')
                            
                            # Studienschwerpunkte
                            schwerpunkte = ausbildung.get('schwerpunkte', '')
                            
                            # Rechte Spalte Inhalte
                            right_column_content = []
                            
                            # Institution/Abschluss formatieren
                            if institution.startswith("Studium"):
                                right_column_content.append(Paragraph(institution, self.custom_styles['Company']))
                            else:
                                right_column_content.append(Paragraph(f"Studium {institution}", self.custom_styles['Company']))
                            
                            if schwerpunkte:
                                right_column_content.append(Paragraph(f"Studienschwerpunkte: {schwerpunkte}", self.custom_styles['Normal']))
                            if abschluss:
                                right_column_content.append(Paragraph(f"{abschluss}", self.custom_styles['Normal']))
                            
                            # Note 
                            note = ausbildung.get('note', '')
                            if note:
                                right_column_content.append(Paragraph(f"Abschlussnote {note}", self.custom_styles['Normal']))
                            
                            # Erstelle zweispaltiges Layout mit mehr Platz für die rechte Spalte
                            data = [[Paragraph(zeitraum, self.custom_styles['Period']), right_column_content[0]]]
                            
                            # Füge weitere Zeilen hinzu
                            for i in range(1, len(right_column_content)):
                                data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[i]])
                            
                            # Tabelle mit definierter Breite (10% links, 75% rechts)
                            col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                            
                            table = Table(data, colWidths=col_widths)
                            table.setStyle(TableStyle([
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                                ('LEFTPADDING', (1, 0), (1, -1), 2*cm),  # Ca. 2cm Einrückung für bessere Lesbarkeit
                            ]))
                            
                            # Wir verpacken die Tabelle und den Spacer in KeepTogether, damit sie nicht über eine Seite verteilt werden
                            entry_elements = [table, Spacer(1, 0.3*cm)]
                            elements.append(KeepTogether(entry_elements))
                        except Exception as e:
                            print(f"Fehler bei der Verarbeitung einer Ausbildung: {str(e)}")
                            # Einfache Darstellung als Fallback
                            elements.append(Paragraph(f"{zeitraum} - {institution}", self.custom_styles['Normal']))
                            elements.append(Spacer(1, 0.3*cm))
                else:
                    elements.append(Paragraph("Keine Ausbildung angegeben", self.custom_styles['Normal']))
                
                # Trennlinie zwischen Ausbildung und Weiterbildungen
                elements.append(Spacer(1, 0.5*cm))
                elements.append(HRFlowable(width="100%", thickness=0.5, lineCap='round', color=colors.lightgrey, spaceBefore=0.3*cm, spaceAfter=0.3*cm))
                
                # Fort- und Weiterbildungen Überschrift
                elements.append(Paragraph("Fort- und Weiterbildungen", self.custom_styles['Heading2']))
                
                # Weiterbildungen im gleichen Format wie Berufserfahrung darstellen
                weiterbildungen = profile_data.get('weiterbildungen', [])
                if weiterbildungen:
                    for weiterbildung in weiterbildungen:
                        try:
                            # Zeitraum
                            zeitraum = weiterbildung.get('zeitraum', '')
                            
                            # Bezeichnung und Abschluss
                            bezeichnung = weiterbildung.get('bezeichnung', '')
                            abschluss = weiterbildung.get('abschluss', '')
                            
                            # Rechte Spalte Inhalte
                            right_column_content = []
                            
                            # Formatieren wie im Design
                            if "zum" in bezeichnung or "zur" in bezeichnung:
                                right_column_content.append(Paragraph(f"Fortbildung {bezeichnung}", self.custom_styles['Company']))
                            else:
                                right_column_content.append(Paragraph(f"Fortbildung zum {bezeichnung}", self.custom_styles['Company']))
                            
                            # Abschluss nur anzeigen, wenn nicht leer und nicht bereits in Bezeichnung enthalten
                            if abschluss and abschluss not in bezeichnung:
                                right_column_content.append(Paragraph(f"{abschluss}", self.custom_styles['Normal']))
                            
                            # Erstelle zweispaltiges Layout mit mehr Platz für die rechte Spalte
                            data = [[Paragraph(zeitraum, self.custom_styles['Period']), right_column_content[0]]]
                            
                            # Füge weitere Zeilen hinzu
                            for i in range(1, len(right_column_content)):
                                data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[i]])
                            
                            # Tabelle mit definierter Breite (10% links, 75% rechts)
                            col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                            
                            table = Table(data, colWidths=col_widths)
                            table.setStyle(TableStyle([
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                                ('LEFTPADDING', (1, 0), (1, -1), 2*cm),  # Ca. 2cm Einrückung für bessere Lesbarkeit
                            ]))
                            
                            # Wir verpacken die Tabelle und den Spacer in KeepTogether, damit sie nicht über eine Seite verteilt werden
                            entry_elements = [table, Spacer(1, 0.3*cm)]
                            elements.append(KeepTogether(entry_elements))
                        except Exception as e:
                            print(f"Fehler bei der Verarbeitung einer Weiterbildung: {str(e)}")
                            # Einfache Darstellung als Fallback
                            elements.append(Paragraph(f"{zeitraum} - {bezeichnung}", self.custom_styles['Normal']))
                            elements.append(Spacer(1, 0.3*cm))
                else:
                    elements.append(Paragraph("Keine Weiterbildungen angegeben", self.custom_styles['Normal']))
                
                # Footer mit GALDORA Kontaktinformationen für beide Templates
                elements.append(Spacer(1, 1.5*cm))
                elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.lightgrey, spaceBefore=0.5*cm))
                elements.append(Spacer(1, 0.2*cm))
                
                footer_text = "GALDORA Personalmanagement GmbH Co.KG\nVolksgartenstr. 85-89, 41065 Mönchengladbach\nE-Mail: info@galdora.de / Web: www.galdora.de"
                elements.append(Paragraph(footer_text, self.custom_styles['Footer']))
                
                return elements
            else:  # Classic Template (default)
                # Ensure we have personal_data initialized at the start
                personal_data = profile_data.get('persönliche_daten', {})
                
                # GALDORA Logo aus dem richtigen Ordner einbinden
                logo_path = get_image_path('Galdoralogo.png', use_static=use_https_compatible)
                
                # Profilbild abrufen, falls vorhanden
                profile_image_path = personal_data.get("profile_image", None)
                profile_img = None
                
                # Versuche, das Profilbild zu laden, wenn es existiert
                if profile_image_path and os.path.exists(profile_image_path):
                    try:
                        # Profilbild laden und anzeigen
                        from PIL import Image as PILImage
                        img_pil = PILImage.open(profile_image_path)
                        img_width, img_height = img_pil.size
                        aspect_ratio = img_width / img_height
                        
                        # Standardgröße für das Profilbild rechts
                        photo_width = 100
                        photo_height = photo_width / aspect_ratio
                        
                        # Profilbild im PDF einfügen
                        profile_img = Image(profile_image_path, width=photo_width, height=photo_height)
                    except Exception as e:
                        print(f"Fehler beim Laden des Profilbilds: {str(e)}")
                        profile_img = None
            
                # Erstelle eine Tabelle für das Logo oben
                if os.path.exists(logo_path) and os.path.isfile(logo_path):
                    try:
                        # Logo-Größe korrigieren (Original-Proportionen beibehalten)
                        from PIL import Image as PILImage
                        img_pil = PILImage.open(logo_path)
                        img_width, img_height = img_pil.size
                        aspect_ratio = img_width / img_height
                        
                        # Anpassung an das Design (kleineres Logo)
                        target_width = 180
                        target_height = target_width / aspect_ratio
                        
                        img = Image(logo_path, width=target_width, height=target_height)
                        
                        # Platzhalter oder Profilbild für die rechte Seite
                        right_content = profile_img if profile_img else Paragraph("", self.custom_styles['Normal'])
                        
                        # Logo-Tabelle mit zwei Spalten: Logo links, Platz für Profil rechts
                        logo_table = Table([[img, right_content]], 
                                    colWidths=[target_width + 10, A4[0] - target_width - 50*mm])
                        
                        logo_table.setStyle(TableStyle([
                            ('ALIGN', (0, 0), (0, 0), 'LEFT'),   # Logo links ausrichten
                            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),  # Profilbild rechts ausrichten
                            ('VALIGN', (0, 0), (1, 0), 'TOP'),   # Beide Zellen oben ausrichten
                        ]))
                        elements.append(logo_table)
                        
                        # Kleinerer Abstand nach dem Logo
                        elements.append(Spacer(1, 0.3*cm))
                    except Exception as e:
                        print(f"Fehler beim Laden des Logos: {str(e)}")
                        # Fallback wenn Bild nicht geladen werden konnte
                        elements.append(Paragraph("GALDORA", self.custom_styles['GaldoraLogo']))
                else:
                    print(f"Logo-Datei nicht gefunden: {logo_path}")
                    # Fallback wenn Bild nicht gefunden wurde
                    elements.append(Paragraph("GALDORA", self.custom_styles['GaldoraLogo']))
            
            # Profil Überschrift - nach links ausgerichtet
            elements.append(Paragraph("Profil", self.custom_styles['ProfilTitle']))
            
            # Persönliche Daten
            personal_data = profile_data.get('persönliche_daten', {})
            
            # Name - nach links ausgerichtet wie gewünscht
            name = personal_data.get('name', 'Profil')
            elements.append(Paragraph(name, self.custom_styles['Name']))
            
            # Ansprechpartner Block mit IHR ANSPRECHPARTNER und den entsprechenden Informationen
            elements.append(Paragraph("IHR ANSPRECHPARTNER", self.custom_styles['ContactHeader']))
            
            # Kontaktinformationen
            kontakt = personal_data.get('kontakt', {})
            ansprechpartner = kontakt.get('ansprechpartner', '')
            
            # Bestimme die korrekte Anrede (Ausnahmebehandlung für Melike Demirkol)
            if ansprechpartner == "Melike Demirkol":
                anrede = f"Frau Demirkol"
            else:
                nachname = ansprechpartner.split()[-1] if ansprechpartner else 'Fischer'
                anrede = f"Herr {nachname}"
            
            telefon = kontakt.get('telefon', '02161 62126-02')
            email = kontakt.get('email', f"{nachname.lower()}@galdora.de")
                
            # Erstelle Layout für Ansprechpartner
            elements.append(Paragraph(anrede, self.custom_styles['ContactData']))
            elements.append(Paragraph(f"{telefon}", self.custom_styles['ContactData']))
            elements.append(Paragraph(f"{email}", self.custom_styles['ContactData']))
            
            # Horizontale Linie nach Ansprechpartner mit etwas Abstand
            elements.append(Spacer(1, 0.5*cm))
            elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.lightgrey))
            elements.append(Spacer(1, 0.5*cm))
                
            # Persönliche Informationen
            elements.append(Paragraph(f"Wohnort: {personal_data.get('wohnort', '')}", self.custom_styles['LabelInline']))
            elements.append(Paragraph(f"Jahrgang: {personal_data.get('jahrgang', '')}", self.custom_styles['LabelInline']))
            elements.append(Paragraph(f"Führerschein: {personal_data.get('führerschein', '')}", self.custom_styles['LabelInline']))
            
            # Wunschgehalt (wenn vorhanden)
            wunschgehalt = profile_data.get('wunschgehalt', '')
            if wunschgehalt:
                elements.append(Paragraph(f"Gehalt: {wunschgehalt}", self.custom_styles['LabelInline']))
            
            # Verfügbarkeit (wenn vorhanden)
            verfuegbarkeit_status = profile_data.get('verfuegbarkeit_status', '')
            if verfuegbarkeit_status:
                elements.append(Paragraph(f"Verfügbarkeit: {verfuegbarkeit_status}", self.custom_styles['LabelInline']))
                
                # Details zur Verfügbarkeit, wenn vorhanden
                verfuegbarkeit_details = profile_data.get('verfuegbarkeit_details', '')
                if verfuegbarkeit_details:
                    elements.append(Paragraph(f"Details zur Verfügbarkeit: {verfuegbarkeit_details}", self.custom_styles['Normal']))
            
            # Zusätzliche Trennlinie vor dem beruflichen Werdegang
            elements.append(Spacer(1, 0.5*cm))
            elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.lightgrey))
            elements.append(Spacer(1, 0.3*cm))
            
            # Beruflicher Werdegang (unterstrichen)
            elements.append(Paragraph("Beruflicher Werdegang", self.custom_styles['Heading2']))
            
            # Berufserfahrung mit einheitlicher Struktur basierend auf dem Design
            berufserfahrung = profile_data.get('berufserfahrung', [])
            if berufserfahrung:
                for erfahrung in berufserfahrung:
                    try:
                        # Zeitraum
                        zeitraum = erfahrung.get('zeitraum', '')
                        
                        # Unternehmen und Position
                        unternehmen = erfahrung.get('unternehmen', '')
                        position = erfahrung.get('position', '')
                        
                        # Aufgaben auf maximal 4 begrenzen
                        aufgaben = erfahrung.get('aufgaben', [])
                        aufgaben_formatted = []
                        for i, aufgabe in enumerate(aufgaben[:4]):  # Maximal 4 Aufgaben
                            aufgaben_formatted.append(Paragraph(f"• {aufgabe}", self.custom_styles['Normal']))
                        
                        # Erstelle zweispaltiges Layout mit mehr Platz für die rechte Spalte
                        # Linke Spalte: Zeitraum
                        # Rechte Spalte: Unternehmen, Position, Aufgaben
                        
                        # Bereite die rechte Spalte vor
                        right_column_content = [
                            Paragraph(unternehmen, self.custom_styles['Company']),
                            Paragraph(position, self.custom_styles['Position'])
                        ]
                        
                        # Aufgaben zur rechten Spalte hinzufügen
                        right_column_content.extend(aufgaben_formatted)
                        
                        # Bereite die Tabellendaten vor
                        data = [[Paragraph(zeitraum, self.custom_styles['Period']), right_column_content[0]]]
                        
                        # Füge weitere Zeilen hinzu
                        for i in range(1, len(right_column_content)):
                            data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[i]])
                        
                        # Tabelle mit definierter Breite (10% links, 75% rechts)
                        col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                        
                        table = Table(data, colWidths=col_widths)
                        table.setStyle(TableStyle([
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                            ('LEFTPADDING', (1, 0), (1, -1), 2*cm),  # Ca. 2cm Einrückung für bessere Lesbarkeit
                        ]))
                        
                        # Wir verpacken die Tabelle und den Spacer in KeepTogether, damit sie nicht über eine Seite verteilt werden
                        entry_elements = [table, Spacer(1, 0.3*cm)]
                        elements.append(KeepTogether(entry_elements))
                    except Exception as e:
                        print(f"Fehler bei der Verarbeitung einer Berufserfahrung: {str(e)}")
                        # Einfache Darstellung als Fallback
                        elements.append(Paragraph(f"{zeitraum} - {unternehmen} - {position}", self.custom_styles['Normal']))
                        elements.append(Spacer(1, 0.3*cm))
            else:
                elements.append(Paragraph("Keine Berufserfahrung angegeben", self.custom_styles['Normal']))
            
            # Ausbildung
            elements.append(Spacer(1, 0.5*cm))
            # Trennlinie zwischen Berufserfahrung und Ausbildung
            elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.lightgrey))
            elements.append(Spacer(1, 0.5*cm))
            elements.append(Paragraph("Ausbildung/ Weiterbildung", self.custom_styles['Heading2']))
            
            # Ausbildungen im gleichen Format wie Berufserfahrung darstellen
            ausbildungen = profile_data.get('ausbildung', [])
            if ausbildungen:
                for ausbildung in ausbildungen:
                    try:
                        # Zeitraum
                        zeitraum = ausbildung.get('zeitraum', '')
                        
                        # Institution und Abschluss
                        institution = ausbildung.get('institution', '')
                        abschluss = ausbildung.get('abschluss', '')
                        
                        # Studienschwerpunkte
                        schwerpunkte = ausbildung.get('schwerpunkte', '')
                        
                        # Rechte Spalte Inhalte
                        right_column_content = []
                        
                        # Institution/Abschluss formatieren
                        if institution.startswith("Studium"):
                            right_column_content.append(Paragraph(institution, self.custom_styles['Company']))
                        else:
                            right_column_content.append(Paragraph(f"Studium {institution}", self.custom_styles['Company']))
                        
                        if schwerpunkte:
                            right_column_content.append(Paragraph(f"Studienschwerpunkte: {schwerpunkte}", self.custom_styles['Normal']))
                        if abschluss:
                            right_column_content.append(Paragraph(f"{abschluss}", self.custom_styles['Normal']))
                        
                        # Note 
                        note = ausbildung.get('note', '')
                        if note:
                            right_column_content.append(Paragraph(f"Abschlussnote {note}", self.custom_styles['Normal']))
                        
                        # Erstelle zweispaltiges Layout mit mehr Platz für die rechte Spalte
                        data = [[Paragraph(zeitraum, self.custom_styles['Period']), right_column_content[0]]]
                        
                        # Füge weitere Zeilen hinzu
                        for i in range(1, len(right_column_content)):
                            data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[i]])
                        
                        # Tabelle mit definierter Breite (10% links, 75% rechts)
                        col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                        
                        table = Table(data, colWidths=col_widths)
                        table.setStyle(TableStyle([
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                            ('LEFTPADDING', (1, 0), (1, -1), 2*cm),  # Ca. 2cm Einrückung für bessere Lesbarkeit
                        ]))
                        
                        # Wir verpacken die Tabelle und den Spacer in KeepTogether, damit sie nicht über eine Seite verteilt werden
                        entry_elements = [table, Spacer(1, 0.3*cm)]
                        elements.append(KeepTogether(entry_elements))
                    except Exception as e:
                        print(f"Fehler bei der Verarbeitung einer Ausbildung: {str(e)}")
                        # Einfache Darstellung als Fallback
                        elements.append(Paragraph(f"{zeitraum} - {institution}", self.custom_styles['Normal']))
                        elements.append(Spacer(1, 0.3*cm))
            else:
                elements.append(Paragraph("Keine Ausbildung angegeben", self.custom_styles['Normal']))
            
            # Trennlinie zwischen Ausbildung und Weiterbildungen
            elements.append(Spacer(1, 0.5*cm))
            elements.append(HRFlowable(width="100%", thickness=0.5, lineCap='round', color=colors.lightgrey, spaceBefore=0.3*cm, spaceAfter=0.3*cm))
            elements.append(Spacer(1, 0.3*cm))
            
            # Fort- und Weiterbildungen Überschrift
            elements.append(Paragraph("Fort- und Weiterbildungen", self.custom_styles['Heading2']))
            
            # Weiterbildungen im gleichen Format wie Berufserfahrung darstellen
            weiterbildungen = profile_data.get('weiterbildungen', [])
            if weiterbildungen:
                for weiterbildung in weiterbildungen:
                    try:
                        # Zeitraum
                        zeitraum = weiterbildung.get('zeitraum', '')
                        
                        # Bezeichnung und Abschluss
                        bezeichnung = weiterbildung.get('bezeichnung', '')
                        abschluss = weiterbildung.get('abschluss', '')
                        
                        # Rechte Spalte Inhalte
                        right_column_content = []
                        
                        # Formatieren wie im Design
                        if "zum" in bezeichnung or "zur" in bezeichnung:
                            right_column_content.append(Paragraph(f"Fortbildung {bezeichnung}", self.custom_styles['Company']))
                        else:
                            right_column_content.append(Paragraph(f"Fortbildung zum {bezeichnung}", self.custom_styles['Company']))
                        
                        # Abschluss nur anzeigen, wenn nicht leer und nicht bereits in Bezeichnung enthalten
                        if abschluss and abschluss not in bezeichnung:
                            right_column_content.append(Paragraph(f"{abschluss}", self.custom_styles['Normal']))
                        
                        # Erstelle zweispaltiges Layout mit mehr Platz für die rechte Spalte
                        data = [[Paragraph(zeitraum, self.custom_styles['Period']), right_column_content[0]]]
                        
                        # Füge weitere Zeilen hinzu
                        for i in range(1, len(right_column_content)):
                            data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[i]])
                        
                        # Tabelle mit definierter Breite (10% links, 75% rechts)
                        col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                        
                        table = Table(data, colWidths=col_widths)
                        table.setStyle(TableStyle([
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                            ('LEFTPADDING', (1, 0), (1, -1), 2*cm),  # Ca. 2cm Einrückung für bessere Lesbarkeit
                        ]))
                        
                        # Wir verpacken die Tabelle und den Spacer in KeepTogether, damit sie nicht über eine Seite verteilt werden
                        entry_elements = [table, Spacer(1, 0.3*cm)]
                        elements.append(KeepTogether(entry_elements))
                    except Exception as e:
                        print(f"Fehler bei der Verarbeitung einer Weiterbildung: {str(e)}")
                        # Einfache Darstellung als Fallback
                        elements.append(Paragraph(f"{zeitraum} - {bezeichnung}", self.custom_styles['Normal']))
                        elements.append(Spacer(1, 0.3*cm))
            else:
                elements.append(Paragraph("Keine Weiterbildungen angegeben", self.custom_styles['Normal']))
            
            # Kein Footer mehr hier, da dieser nun über das PageTemplate eingefügt wird
            
            return elements
        except Exception as e:
            print(f"Fehler bei der Erstellung der Dokumentelemente: {str(e)}")
            # Einfache Elemente als Fallback
            elements.append(Paragraph("FEHLER BEI DER ERSTELLUNG DES PROFILS", self.styles['Title']))
            elements.append(Paragraph(f"Details: {str(e)}", self.styles['Normal']))
        return elements


def check_missing_profile_data(profile_data):
    """
    Überprüft, ob wichtige Daten im Profil fehlen
    
    Args:
        profile_data: Dictionary mit Profildaten
    
    Returns:
        Liste mit fehlenden Datenfeldern
    """
    missing_data = []
    
    # Persönliche Daten prüfen
    personal_data = profile_data.get('persönliche_daten', {})
    if not personal_data.get('name'):
        missing_data.append('Name')
    
    # Kontaktdaten prüfen
    kontakt = personal_data.get('kontakt', {})
    if not kontakt.get('email') and not kontakt.get('telefon'):
        missing_data.append('Kontaktdaten (E-Mail oder Telefon)')
    
    # Wohnort prüfen
    if not personal_data.get('wohnort'):
        missing_data.append('Wohnort')
    
    # Jahrgang prüfen
    if not personal_data.get('jahrgang'):
        missing_data.append('Jahrgang')
    
    # Berufserfahrung prüfen
    berufserfahrung = profile_data.get('berufserfahrung', [])
    if not berufserfahrung:
        missing_data.append('Berufserfahrung')
    
    # Ausbildung prüfen
    ausbildung = profile_data.get('ausbildung', [])
    if not ausbildung:
        missing_data.append('Ausbildung')
    
    return missing_data
